from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


_TITLE_RE = re.compile(r"^\s*#{1,6}\s+(.*)$")
_LIST_RE = re.compile(r"^\s*(?:[-*•]|\d+\.)\s+(.*)$")


def txt_to_docx(
    txt_path: str | Path,
    docx_path: str | Path,
    *,
    encoding: str = "utf-8",
    base_font: str = "Times New Roman",
    base_size_pt: int = 11,
) -> Path:
    """
    Convertit un fichier TXT en DOCX structuré.

    Règles simples :
    - lignes vides → séparation de paragraphes
    - lignes '# Titre' ou MAJUSCULES → titres
    - listes (-, *, •, 1.) → listes Word
    """

    txt_path = Path(txt_path)
    docx_path = Path(docx_path)

    if not txt_path.exists():
        raise FileNotFoundError(txt_path)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = base_font
    style.font.size = Pt(base_size_pt)

    with open(txt_path, "r", encoding=encoding) as f:
        for raw_line in f:
            line = raw_line.rstrip()

            if not line:
                doc.add_paragraph("")
                continue

            # ───────────────
            # Titres Markdown
            # ───────────────
            m = _TITLE_RE.match(line)
            if m:
                title = m.group(1).strip()
                p = doc.add_heading(title, level=1)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue

            # ───────────────
            # Titres en MAJUSCULES
            # ───────────────
            if line.isupper() and len(line) > 3:
                p = doc.add_heading(line, level=2)
                continue

            # ───────────────
            # Listes
            # ───────────────
            m = _LIST_RE.match(line)
            if m:
                item = m.group(1).strip()
                doc.add_paragraph(item, style="List Bullet")
                continue

            # ───────────────
            # Paragraphe normal
            # ───────────────
            doc.add_paragraph(line)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    return docx_path


__all__ = ["txt_to_docx"]

from __future__ import annotations

from pathlib import Path
from typing import Optional, Iterable

import csv

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
except Exception:  # pragma: no cover
    Document = None


class CSV2DocxError(RuntimeError):
    pass


def _require_docx():
    if Document is None:
        raise CSV2DocxError(
            "python-docx non installé. Installez-le avec: pip install python-docx"
        )


def _normalize_row(row: Iterable[str], width: int) -> list[str]:
    row = list(row)
    if len(row) < width:
        row += [""] * (width - len(row))
    elif len(row) > width:
        row = row[:width]
    return [str(c) for c in row]


def convert(
    csv_path: str | Path,
    docx_path: str | Path,
    *,
    delimiter: str = ",",
    encoding: str = "utf-8",
    has_header: Optional[bool] = None,
    font_size_pt: int = 10,
    table_style: str = "Table Grid",
) -> Path:
    """
    Convertit un CSV en DOCX via python-docx.

    - streaming CSV (pas de chargement complet en mémoire)
    - support BOM UTF-8
    - normalisation largeur lignes
    - style de table configurable
    """

    _require_docx()

    csv_path = Path(csv_path)
    docx_path = Path(docx_path)

    if not csv_path.exists():
        raise FileNotFoundError(csv_path)

    doc = Document()
    doc.add_heading(csv_path.stem, level=1)

    try:
        f = open(csv_path, newline="", encoding=encoding)
    except UnicodeDecodeError:
        f = open(csv_path, newline="", encoding="utf-8-sig")

    with f:
        reader = csv.reader(f, delimiter=delimiter)

        try:
            first_row = next(reader)
        except StopIteration:
            doc.add_paragraph("")
            docx_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(docx_path)
            return docx_path

        # Détection simple de header si non précisé
        if has_header is None:
            has_header = any(not cell.replace(".", "").isdigit() for cell in first_row)

        rows_iter = reader if has_header else iter([first_row, *reader])
        header = first_row if has_header else None

        width = max(len(first_row), *(len(r) for r in rows_iter))
        rows_iter = reader if has_header else iter([first_row, *reader])

        table = doc.add_table(rows=0, cols=width)
        table.style = table_style

        def _write_row(values: Iterable[str], bold: bool = False) -> None:
            cells = table.add_row().cells
            for i, val in enumerate(_normalize_row(values, width)):
                p = cells[i].paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(font_size_pt)
                run.bold = bold

        if header:
            _write_row(header, bold=True)

        for row in rows_iter:
            _write_row(row, bold=False)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    return docx_path


__all__ = ["convert", "CSV2DocxError"]

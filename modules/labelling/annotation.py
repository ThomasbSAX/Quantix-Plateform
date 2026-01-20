from __future__ import annotations

"""Backend de labelling (inspiré Label Studio).

Le module expose:
- une API FastAPI optionnelle (si `fastapi` + `pydantic` sont installés),
- un blueprint Flask via `create_flask_blueprint()` (recommandé pour ce projet).

Contraintes:
- import-safe côté Flask: aucune dépendance lourde obligatoire (FastAPI, torch/transformers, etc.).
"""

import csv
import io
import json
import os
import re
import tempfile
import threading
import uuid
import zipfile
from dataclasses import dataclass, field
from datetime import datetime
from functools import lru_cache
from typing import Any, Dict, Iterable, List, Optional, Pattern, Sequence, Tuple

# ---------------------------
# Dépendances optionnelles
# ---------------------------

try:
    from fastapi import FastAPI, File, Form, Query, UploadFile
    from fastapi.middleware.cors import CORSMiddleware
    from fastapi.responses import JSONResponse, StreamingResponse
except Exception:  # pragma: no cover
    FastAPI = None
    File = None
    Form = None
    Query = None
    UploadFile = Any  # type: ignore[assignment]
    CORSMiddleware = None
    JSONResponse = None
    StreamingResponse = None

try:
    from pydantic import BaseModel
except Exception:  # pragma: no cover
    BaseModel = None

try:
    import pdfplumber
except Exception:  # pragma: no cover
    pdfplumber = None

try:
    import docx  # python-docx
except Exception:  # pragma: no cover
    docx = None

try:
    import torch
except Exception:  # pragma: no cover
    torch = None

try:
    from transformers import pipeline
except Exception:  # pragma: no cover
    pipeline = None

try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter
except Exception:  # pragma: no cover
    Workbook = None
    Alignment = None
    Font = None
    PatternFill = None
    get_column_letter = None


@dataclass
class Span:
    start: int
    end: int


@dataclass
class Document:
    id: str
    name: str
    raw_text: str
    created_at: str
    meta: Dict[str, Any] = field(default_factory=dict)


@dataclass
class Label:
    id: str
    name: str
    color_hex: str


@dataclass
class Annotation:
    id: str
    doc_id: str
    span: Span
    labels: List[str]
    weight: float
    created_at: str
    meta: Dict[str, Any] = field(default_factory=dict)
    text: str = ""
    sentence: str = ""


@dataclass
class Relation:
    id: str
    doc_id: str
    source_ann_id: str
    target_ann_id: str
    rel_type: str
    weight: float
    created_at: str
    meta: Dict[str, Any] = field(default_factory=dict)
# Utilities: sentence slicing + validation
# ---------------------------

_SENT_SPLIT = re.compile(r"(?<=[.!?])\s+|\n+")

def split_sentences(text: str) -> List[Tuple[int, int, str]]:
    """
    Returns list of (start_idx, end_idx, sentence_text).
    Sentence boundaries are heuristic; UI may override by sending the sentence range directly.
    """
    out: List[Tuple[int, int, str]] = []
    if not text:
        return out
    idx = 0
    for part in _SENT_SPLIT.split(text):
        if not part:
            idx += 1
            continue
        start = text.find(part, idx)
        if start < 0:
            start = idx
        end = start + len(part)
        out.append((start, end, part))
        idx = end
    return out

def sentence_for_span(text: str, span: Span, precomputed: Optional[List[Tuple[int,int,str]]] = None) -> str:
    sents = precomputed if precomputed is not None else split_sentences(text)
    if not sents:
        return ""
    # find sentence that overlaps the span (max overlap)
    best = ""
    best_overlap = -1
    for s0, s1, s in sents:
        overlap = max(0, min(span.end, s1) - max(span.start, s0))
        if overlap > best_overlap:
            best_overlap = overlap
            best = s
    return best

def clamp01(x: float) -> float:
    return 0.0 if x < 0.0 else 1.0 if x > 1.0 else float(x)

def now_iso() -> str:
    return datetime.utcnow().replace(microsecond=0).isoformat() + "Z"

def new_id(prefix: str) -> str:
    return f"{prefix}_{uuid.uuid4().hex[:12]}"


# ---------------------------
# Store: in-memory registry (swap with DB later)
# ---------------------------

class AnnotationStore:
    """
    Minimal store: documents, labels, annotations, relations.
    Plug this into an API layer (FastAPI/Flask) and let the UI send spans + label_ids + weights.
    """

    def __init__(self) -> None:
        self.documents: Dict[str, Document] = {}
        self.labels: Dict[str, Label] = {}
        self.annotations: Dict[str, Annotation] = {}
        self.relations: Dict[str, Relation] = {}

    # ---- documents

    def add_document(self, name: str, raw_text: str, *, meta: Optional[Dict[str, Any]] = None) -> Document:
        doc = Document(id=new_id("doc"), name=name, raw_text=raw_text, created_at=now_iso(), meta=meta or {})
        self.documents[doc.id] = doc
        return doc

    def get_document(self, doc_id: str) -> Document:
        if doc_id not in self.documents:
            raise KeyError(f"Unknown doc_id={doc_id}")
        return self.documents[doc_id]

    # ---- labels

    def add_label(self, name: str, color_hex: str) -> Label:
        lab = Label(id=new_id("lab"), name=name, color_hex=color_hex)
        self.labels[lab.id] = lab
        return lab

    def get_label(self, label_id: str) -> Label:
        if label_id not in self.labels:
            raise KeyError(f"Unknown label_id={label_id}")
        return self.labels[label_id]

    # ---- annotations

    def add_annotation(
        self,
        doc_id: str,
        start: int,
        end: int,
        label_ids: Sequence[str],
        weight: float,
        *,
        meta: Optional[Dict[str, Any]] = None,
        sentence_hint: Optional[Tuple[int, int]] = None,  # (sent_start, sent_end) from UI if available
    ) -> Annotation:
        doc = self.get_document(doc_id)

        span = Span(start=start, end=end)
        if span.end > len(doc.raw_text):
            raise ValueError("Span exceeds document text length")
        if not label_ids:
            raise ValueError("At least one label is required")
        for lid in label_ids:
            _ = self.get_label(lid)

        text = doc.raw_text[span.start:span.end]

        if sentence_hint is not None:
            s0, s1 = sentence_hint
            if not (0 <= s0 <= s1 <= len(doc.raw_text)):
                raise ValueError("Invalid sentence_hint range")
            sent = doc.raw_text[s0:s1]
        else:
            sent = sentence_for_span(doc.raw_text, span)

        ann = Annotation(
            id=new_id("ann"),
            doc_id=doc_id,
            span=span,
            labels=list(label_ids),
            weight=clamp01(float(weight)),
            created_at=now_iso(),
            meta=meta or {},
            text=text,
            sentence=sent,
        )
        self.annotations[ann.id] = ann
        return ann

    def update_annotation(
        self,
        ann_id: str,
        *,
        label_ids: Optional[Sequence[str]] = None,
        weight: Optional[float] = None,
        meta_patch: Optional[Dict[str, Any]] = None,
    ) -> Annotation:
        if ann_id not in self.annotations:
            raise KeyError(f"Unknown ann_id={ann_id}")
        ann = self.annotations[ann_id]

        if label_ids is not None:
            if not label_ids:
                raise ValueError("At least one label is required")
            for lid in label_ids:
                _ = self.get_label(lid)
            ann.labels = list(label_ids)

        if weight is not None:
            ann.weight = clamp01(float(weight))

        if meta_patch:
            ann.meta.update(meta_patch)

        return ann

    def list_annotations(self, doc_id: Optional[str] = None) -> List[Annotation]:
        if doc_id is None:
            return list(self.annotations.values())
        return [a for a in self.annotations.values() if a.doc_id == doc_id]

    # ---- relations

    def add_relation(
        self,
        doc_id: str,
        source_ann_id: str,
        target_ann_id: str,
        rel_type: str,
        weight: float,
        *,
        meta: Optional[Dict[str, Any]] = None,
    ) -> Relation:
        _ = self.get_document(doc_id)

        if source_ann_id not in self.annotations:
            raise KeyError(f"Unknown source_ann_id={source_ann_id}")
        if target_ann_id not in self.annotations:
            raise KeyError(f"Unknown target_ann_id={target_ann_id}")

        if self.annotations[source_ann_id].doc_id != doc_id or self.annotations[target_ann_id].doc_id != doc_id:
            raise ValueError("Relation endpoints must belong to the same doc_id")

        rel = Relation(
            id=new_id("rel"),
            doc_id=doc_id,
            source_ann_id=source_ann_id,
            target_ann_id=target_ann_id,
            rel_type=rel_type,
            weight=clamp01(float(weight)),
            created_at=now_iso(),
            meta=meta or {},
        )
        self.relations[rel.id] = rel
        return rel

    def list_relations(self, doc_id: Optional[str] = None) -> List[Relation]:
        if doc_id is None:
            return list(self.relations.values())
        return [r for r in self.relations.values() if r.doc_id == doc_id]


# ---------------------------
# Weight normalization policies
# ---------------------------

class WeightNormalizer:
    """
    Normalize weights to [0,1] per document (or globally).
    Policies:
      - "identity": leave as-is
      - "minmax": (w - min) / (max - min) if max>min else 1 for all nonzero else 0
      - "sum1": w / sum(w) then optionally rescale to [0,1] by dividing by max
    """

    @staticmethod
    def normalize(weights: Sequence[float], policy: str = "identity") -> List[float]:
        ws = [clamp01(float(w)) for w in weights]
        if not ws:
            return []

        policy = (policy or "identity").lower().strip()
        if policy == "identity":
            return ws

        if policy == "minmax":
            w_min = min(ws)
            w_max = max(ws)
            if w_max == w_min:
                # Convention: si tout est égal, on renvoie 0 si tout est 0, sinon 1
                return [0.0 if w_max == 0.0 else 1.0 for _ in ws]
            return [clamp01((w - w_min) / (w_max - w_min)) for w in ws]

        if policy == "sum1":
            s = sum(ws)
            if s == 0.0:
                return [0.0 for _ in ws]
            frac = [w / s for w in ws]
            m = max(frac)
            if m == 0.0:
                return [0.0 for _ in ws]
            return [clamp01(w / m) for w in frac]

        raise ValueError(f"Unknown normalize policy: {policy}")


# ---------------------------
# Service thread-safe + (optionnel) persistance JSON
# ---------------------------

_DEFAULT_DOC_ID = "doc_default"


_BASE_MARKERS: List[Dict[str, str]] = [
    {"name": "POSITIVE", "color": "#00FF00"},
    {"name": "NEGATIVE", "color": "#FF0000"},
    {"name": "NEUTRAL", "color": "#0000FF"},
]


class LabellingService:
    """État + opérations d’annotation.

    Objectif: isoler la logique métier des frameworks web.
    - Thread-safe: verrou interne
    - Persistance optionnelle: JSON (écriture atomique)
    """

    def __init__(self, *, data_path: Optional[str] = None) -> None:
        self._lock = threading.RLock()
        self._data_path = data_path
        self._docs: Dict[str, Dict[str, Any]] = {}
        self._ensure_doc(_DEFAULT_DOC_ID)
        self._load_if_any()

    def _ensure_doc(self, doc_id: str, *, name: Optional[str] = None) -> None:
        if doc_id not in self._docs:
            self._docs[doc_id] = {
                "name": name or doc_id,
                "text": "",
                "annotations": [],
                "relations": [],
                "history": [],
                "history_index": -1,
                "meta": {},
                "markers": [dict(m) for m in _BASE_MARKERS],
                "updated_at": now_iso(),
            }
        else:
            if name is not None:
                self._docs[doc_id]["name"] = name

            # rétro-compat: documents plus anciens sans champs "markers"
            if "markers" not in self._docs[doc_id] or not isinstance(self._docs[doc_id].get("markers"), list):
                self._docs[doc_id]["markers"] = [dict(m) for m in _BASE_MARKERS]

    def _validate_doc_id(self, doc_id: str) -> str:
        doc_id = (doc_id or "").strip()
        if not doc_id:
            raise ValueError("doc_id invalide")
        # on évite des ids avec des chemins/espaces (sécurité basique)
        if any(ch in doc_id for ch in ["/", "\\", " ", "\t", "\n"]):
            raise ValueError("doc_id invalide")
        return doc_id

    def create_document(
        self,
        *,
        name: str,
        text: str = "",
        doc_id: Optional[str] = None,
        meta: Optional[Dict[str, Any]] = None,
    ) -> str:
        with self._lock:
            if doc_id is None:
                doc_id = new_id("doc")
            doc_id = self._validate_doc_id(doc_id)
            if doc_id in self._docs:
                raise ValueError("doc_id déjà existant")
            self._ensure_doc(doc_id, name=name)
            self._docs[doc_id]["text"] = text or ""
            self._docs[doc_id]["meta"] = meta or {}
            self._docs[doc_id]["annotations"] = []
            self._docs[doc_id]["history"] = []
            self._docs[doc_id]["history_index"] = -1
            self._snapshot(doc_id)
            self._docs[doc_id]["updated_at"] = now_iso()
            self._persist()
            return doc_id

    def delete_document(self, *, doc_id: str) -> None:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            if doc_id == _DEFAULT_DOC_ID:
                raise ValueError("Impossible de supprimer le document par défaut")
            if doc_id not in self._docs:
                raise KeyError("Document non trouvé")
            del self._docs[doc_id]
            self._persist()

    def list_documents(self) -> List[Dict[str, Any]]:
        with self._lock:
            out: List[Dict[str, Any]] = []
            for doc_id, doc in self._docs.items():
                out.append(
                    {
                        "doc_id": doc_id,
                        "name": doc.get("name", doc_id),
                        "text_length": len(doc.get("text", "") or ""),
                        "annotation_count": len(doc.get("annotations", []) or []),
                        "updated_at": doc.get("updated_at"),
                        "meta": doc.get("meta", {}),
                    }
                )
            # tri stable: défaut en premier puis date
            out.sort(key=lambda d: (0 if d["doc_id"] == _DEFAULT_DOC_ID else 1, d.get("updated_at") or ""), reverse=False)
            return out

    def get_document_info(self, *, doc_id: str) -> Dict[str, Any]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            doc = self._docs[doc_id]
            return {
                "doc_id": doc_id,
                "name": doc.get("name", doc_id),
                "text_length": len(doc.get("text", "") or ""),
                "annotation_count": len(doc.get("annotations", []) or []),
                "updated_at": doc.get("updated_at"),
                "meta": doc.get("meta", {}),
            }

    def _snapshot(self, doc_id: str) -> None:
        doc = self._docs[doc_id]
        snap = {
            "annotations": list(doc["annotations"]),
            "relations": list(doc.get("relations", [])),
            "timestamp": datetime.utcnow().isoformat(),
        }
        # overwrite future if user had undone
        doc["history"] = doc["history"][: doc["history_index"] + 1]
        doc["history"].append(snap)
        doc["history_index"] = len(doc["history"]) - 1

    def _dump_state(self) -> Dict[str, Any]:
        return {
            "version": 3,
            # Backward-compat: ancien format avec markers top-level
            "markers": _BASE_MARKERS,
            "docs": self._docs,
        }

    def _atomic_write(self, path: str, data: str) -> None:
        d = os.path.dirname(path)
        if d:
            os.makedirs(d, exist_ok=True)
        with tempfile.NamedTemporaryFile("w", encoding="utf-8", delete=False, dir=d or None) as tmp:
            tmp.write(data)
            tmp.flush()
            os.fsync(tmp.fileno())
            tmp_path = tmp.name
        os.replace(tmp_path, path)

    def _persist(self) -> None:
        if not self._data_path:
            return
        payload = json.dumps(self._dump_state(), ensure_ascii=False, indent=2)
        self._atomic_write(self._data_path, payload)

    def _load_if_any(self) -> None:
        if not self._data_path:
            return
        if not os.path.exists(self._data_path):
            return
        try:
            with open(self._data_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict) and "docs" in data:
                self._docs = data.get("docs", {}) or {}
                self._ensure_doc(_DEFAULT_DOC_ID)

                # migration: on s'assure que chaque doc a ses markers
                for doc_id in list(self._docs.keys()):
                    self._ensure_doc(doc_id)

                # migration legacy: si le JSON avait des markers globaux, on les associe
                # uniquement au doc par défaut (sinon nouveaux documents hériteraient).
                legacy_markers = data.get("markers")
                if isinstance(legacy_markers, list) and _DEFAULT_DOC_ID in self._docs:
                    cur = self._docs[_DEFAULT_DOC_ID].get("markers")
                    if not isinstance(cur, list) or cur == _BASE_MARKERS:
                        self._docs[_DEFAULT_DOC_ID]["markers"] = list(legacy_markers)
        except Exception:
            # On ne casse pas le service si le fichier est corrompu
            self._ensure_doc(_DEFAULT_DOC_ID)

    # ---- markers

    def list_markers(self, *, doc_id: str = _DEFAULT_DOC_ID) -> List[Dict[str, str]]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            return list(self._docs[doc_id].get("markers", []) or [])

    def add_marker(self, name: str, color: str, *, doc_id: str = _DEFAULT_DOC_ID) -> List[Dict[str, str]]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            markers = list(self._docs[doc_id].get("markers", []) or [])
            for m in markers:
                if m["name"].lower() == name.lower():
                    raise ValueError("Marqueur déjà existant")
            markers.append({"name": name, "color": color})
            self._docs[doc_id]["markers"] = markers
            self._persist()
            return list(markers)

    def update_marker(self, name: str, color: str, *, doc_id: str = _DEFAULT_DOC_ID) -> Dict[str, str]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            markers = self._docs[doc_id].get("markers", []) or []
            for m in markers:
                if m["name"].lower() == name.lower():
                    m["color"] = color
                    self._docs[doc_id]["markers"] = list(markers)
                    self._persist()
                    return dict(m)
            raise KeyError("Marqueur non trouvé")

    def delete_marker(self, name: str, *, doc_id: str = _DEFAULT_DOC_ID) -> List[Dict[str, str]]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            markers = list(self._docs[doc_id].get("markers", []) or [])
            markers = [m for m in markers if m["name"].lower() != name.lower()]
            self._docs[doc_id]["markers"] = markers
            self._persist()
            return list(markers)

    # ---- documents

    def set_text(self, text: str, *, doc_id: str = _DEFAULT_DOC_ID) -> str:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            self._docs[doc_id]["text"] = text or ""
            self._docs[doc_id]["annotations"] = []
            self._docs[doc_id]["relations"] = []
            self._docs[doc_id]["history"] = []
            self._docs[doc_id]["history_index"] = -1
            self._snapshot(doc_id)
            self._docs[doc_id]["updated_at"] = now_iso()
            self._persist()
            return self._docs[doc_id]["text"]

    def get_text(self, *, doc_id: str = _DEFAULT_DOC_ID) -> str:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            return self._docs[doc_id]["text"]

    # ---- annotations

    def list_annotations(self, *, doc_id: str = _DEFAULT_DOC_ID) -> List[Dict[str, Any]]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            return list(self._docs[doc_id]["annotations"])

    def delete_annotations_by_span(self, *, doc_id: str = _DEFAULT_DOC_ID, span_start: int, span_end: int) -> Dict[str, Any]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            text = self._docs[doc_id].get("text", "") or ""

            if not isinstance(span_start, int) or not isinstance(span_end, int):
                raise ValueError("span_start/span_end invalides")
            if span_start < 0 or span_end <= span_start:
                raise ValueError("Plage invalide")
            if text and span_start > len(text):
                raise ValueError("Plage hors texte")
            if text and span_end > len(text):
                span_end = len(text)

            anns = list(self._docs[doc_id].get("annotations", []) or [])
            kept: List[Dict[str, Any]] = []
            removed: List[Dict[str, Any]] = []

            for a in anns:
                a0 = a.get("span_start")
                a1 = a.get("span_end")
                if isinstance(a0, int) and isinstance(a1, int) and a0 < a1:
                    overlaps = (a0 < span_end) and (a1 > span_start)
                    if overlaps:
                        removed.append(a)
                        continue
                kept.append(a)

            self._docs[doc_id]["annotations"] = kept
            self._snapshot(doc_id)
            self._docs[doc_id]["updated_at"] = now_iso()
            self._persist()
            return {"removed": len(removed), "annotations": list(kept)}

    def _ann_index_by_id(self, *, doc_id: str) -> Dict[str, Dict[str, Any]]:
        doc = self._docs[doc_id]
        out: Dict[str, Dict[str, Any]] = {}
        for a in doc.get("annotations", []) or []:
            ann_id = a.get("ann_id")
            if isinstance(ann_id, str):
                out[ann_id] = a
        return out

    def _infer_span_for_manual(self, *, doc_id: str, mot: str, phrase: str) -> Tuple[Optional[int], Optional[int]]:
        """Infère au mieux span_start/span_end pour une annotation manuelle.

        Si impossible/ambigu: renvoie (None, None).
        """
        text = self._docs[doc_id].get("text", "") or ""
        if not text or not mot:
            return (None, None)

        # 1) si la phrase est un extrait du texte, chercher dans cette zone
        if phrase:
            idx = text.find(phrase)
            if idx >= 0:
                local = phrase
                # recherche insensitive de mot dans la phrase
                m = re.search(re.escape(mot), local, flags=re.IGNORECASE | re.UNICODE)
                if m:
                    return (idx + m.start(), idx + m.end())

        # 2) fallback: première occurrence du mot
        m2 = re.search(re.escape(mot), text, flags=re.IGNORECASE | re.UNICODE)
        if m2:
            return (m2.start(), m2.end())

        return (None, None)

    def _annotation_signature(self, ann: Dict[str, Any]) -> Tuple[Any, ...]:
        # signature stable pour déduplication (ne dépend pas de ann_id/created_at/meta)
        return (
            ann.get("mot"),
            ann.get("marqueur"),
            ann.get("couleur"),
            ann.get("span_start"),
            ann.get("span_end"),
            ann.get("phrase"),
            ann.get("index_phrase"),
        )

    def search_annotations(
        self,
        *,
        doc_id: str = _DEFAULT_DOC_ID,
        mot: Optional[str] = None,
        marqueur: Optional[str] = None,
        phrase: Optional[str] = None,
    ) -> List[Dict[str, Any]]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            anns = self.list_annotations(doc_id=doc_id)
        results = anns
        if mot:
            results = [a for a in results if mot.lower() in str(a.get("mot", "")).lower()]
        if marqueur:
            results = [a for a in results if marqueur.lower() == str(a.get("marqueur", "")).lower()]
        if phrase:
            results = [a for a in results if phrase.lower() in str(a.get("phrase", "")).lower()]
        return results

    def add_annotation(
        self,
        ann: Dict[str, Any],
        *,
        doc_id: str = _DEFAULT_DOC_ID,
        auto_expand: bool = True,
        dedupe: bool = True,
    ) -> int:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            text = self._docs[doc_id]["text"]
            if not text:
                raise ValueError("Aucun texte chargé")

            required = ["mot", "marqueur", "couleur", "phrase", "index_phrase"]
            for k in required:
                if k not in ann:
                    raise ValueError(f"Champ manquant: {k}")

            # insertion principale (annotation manuelle)
            manual_ann = dict(ann)
            manual_ann.setdefault("ann_id", new_id("ann"))
            manual_ann.setdefault("doc_id", doc_id)
            manual_ann.setdefault("created_at", now_iso())
            manual_ann.setdefault("meta", {})
            if isinstance(manual_ann.get("meta"), dict):
                manual_ann["meta"].setdefault("source", "manual")
                manual_ann["meta"].setdefault("created_at", now_iso())

            # spans: acceptés si fournis par le frontend, sinon tentative d'inférence
            span_start = manual_ann.get("span_start")
            span_end = manual_ann.get("span_end")
            if not (isinstance(span_start, int) and isinstance(span_end, int) and 0 <= span_start < span_end <= len(text)):
                s0, s1 = self._infer_span_for_manual(
                    doc_id=doc_id,
                    mot=str(manual_ann.get("mot", "")),
                    phrase=str(manual_ann.get("phrase", "")),
                )
                manual_ann["span_start"] = s0
                manual_ann["span_end"] = s1

            # contexte (utile entraînement)
            if isinstance(manual_ann.get("span_start"), int) and isinstance(manual_ann.get("span_end"), int):
                s = max(0, int(manual_ann["span_start"]) - 80)
                e = min(len(text), int(manual_ann["span_end"]) + 80)
                manual_ann.setdefault("context_start", s)
                manual_ann.setdefault("context_end", e)
                manual_ann.setdefault("context", text[s:e])

            if dedupe:
                sig = self._annotation_signature(manual_ann)
                existing = {self._annotation_signature(a) for a in self._docs[doc_id]["annotations"]}
                if sig in existing:
                    # déjà présent: on ne recrée pas
                    if auto_expand:
                        self._annotate_similar_words(
                            doc_id=doc_id,
                            text=text,
                            mot=str(ann["mot"]),
                            marqueur=str(ann["marqueur"]),
                            couleur=str(ann["couleur"]),
                            meta_base={
                                "source": "auto_regex",
                                "seed": {"mot": str(ann["mot"]), "marqueur": str(ann["marqueur"])},
                                "created_at": now_iso(),
                            },
                        )
                    self._snapshot(doc_id)
                    self._docs[doc_id]["updated_at"] = now_iso()
                    self._persist()
                    return len(self._docs[doc_id]["annotations"])

                self._docs[doc_id]["annotations"].append(manual_ann)
            else:
                self._docs[doc_id]["annotations"].append(manual_ann)

            if auto_expand:
                self._annotate_similar_words(
                    doc_id=doc_id,
                    text=text,
                    mot=str(ann["mot"]),
                    marqueur=str(ann["marqueur"]),
                    couleur=str(ann["couleur"]),
                    meta_base={
                        "source": "auto_regex",
                        "seed": {"mot": str(ann["mot"]), "marqueur": str(ann["marqueur"])},
                        "created_at": now_iso(),
                    },
                )

            self._snapshot(doc_id)
            self._docs[doc_id]["updated_at"] = now_iso()
            self._persist()
            return len(self._docs[doc_id]["annotations"])

    def _annotate_similar_words(
        self,
        *,
        doc_id: str,
        text: str,
        mot: str,
        marqueur: str,
        couleur: str,
        max_matches: int = 2000,
        meta_base: Optional[Dict[str, Any]] = None,
    ) -> None:
        # Variante robuste: tolère accents/français + pluriels simples.
        # On limite le nombre de matches pour éviter les explosions sur des patterns trop larges.
        vowels = "aeiouyàâäæéèêëîïôöùûüÿ"
        mot_escaped = re.escape(mot.strip())
        if not mot_escaped:
            return

        # Remplace chaque voyelle par une classe "voyelles" (très approximatif mais pratique)
        def _vowelize(s: str) -> str:
            out = []
            for ch in s:
                if ch.lower() in vowels:
                    out.append(f"[{vowels}{vowels.upper()}]")
                else:
                    out.append(ch)
            return "".join(out)

        mot_regex = _vowelize(mot_escaped)
        # Support des mentions multi-mots: espace(s) => \s+
        mot_regex = mot_regex.replace("\\ ", r"\\s+")
        # Si multi-mots, on encadre par des limites souples, sinon frontières de mots
        if r"\\s+" in mot_regex:
            pattern = rf"(?<!\w){mot_regex}(?!\w)"
        else:
            pattern = rf"\b{mot_regex}\w*\b"
        matches = list(re.finditer(pattern, text, flags=re.IGNORECASE | re.UNICODE))
        if len(matches) > max_matches:
            matches = matches[:max_matches]

        phrases = re.split(r"\.(?:\s+|$)", text)
        for match in matches:
            found = match.group()
            # Cherche une phrase contenant le mot; fallback sur un extrait local.
            phrase_txt = ""
            idx_phrase = 0
            for idx, phr in enumerate(phrases):
                if found.lower() in phr.lower():
                    phrase_txt = phr.strip()
                    idx_phrase = idx
                    break
            if not phrase_txt:
                s = max(0, match.start() - 80)
                e = min(len(text), match.end() + 80)
                phrase_txt = text[s:e].strip()

            candidate = {
                "ann_id": new_id("ann"),
                "doc_id": doc_id,
                "mot": found,
                "marqueur": marqueur,
                "couleur": couleur,
                "phrase": phrase_txt,
                "index_phrase": idx_phrase,
                "span_start": int(match.start()),
                "span_end": int(match.end()),
                "context_start": max(0, match.start() - 80),
                "context_end": min(len(text), match.end() + 80),
                "context": text[max(0, match.start() - 80) : min(len(text), match.end() + 80)],
                "created_at": now_iso(),
            }
            if meta_base is not None:
                candidate["meta"] = dict(meta_base)
            if "meta" not in candidate:
                candidate["meta"] = {"source": "auto_regex", "created_at": now_iso()}

            if candidate.get("span_start") == candidate.get("span_end"):
                continue

            # dédup: par signature
            sig = self._annotation_signature(candidate)
            existing = {self._annotation_signature(a) for a in self._docs[doc_id]["annotations"]}
            if sig not in existing:
                self._docs[doc_id]["annotations"].append(candidate)

    def propagate_annotation_to_docs(
        self,
        *,
        source_doc_id: str,
        target_doc_ids: Sequence[str],
        annotation: Dict[str, Any],
        include_source_doc: bool = False,
        max_matches_per_doc: int = 2000,
    ) -> Dict[str, Any]:
        """Propage une annotation "seed" sur plusieurs documents.

        Cas d’usage: l’utilisateur annote manuellement "Napoléon" sur un doc,
        puis veut appliquer automatiquement la même annotation (regex) sur d’autres docs sélectionnés.
        """

        with self._lock:
            source_doc_id = self._validate_doc_id(source_doc_id)
            self._ensure_doc(source_doc_id)

            op_id = new_id("prop")
            mot = str(annotation.get("mot", "")).strip()
            marqueur = str(annotation.get("marqueur", "")).strip()
            couleur = str(annotation.get("couleur", "")).strip()
            if not mot or not marqueur or not couleur:
                raise ValueError("annotation doit contenir mot/marqueur/couleur")

            targets = [self._validate_doc_id(d) for d in target_doc_ids]
            if include_source_doc and source_doc_id not in targets:
                targets.append(source_doc_id)

            results: Dict[str, Any] = {"op_id": op_id, "source_doc_id": source_doc_id, "targets": {}, "skipped": []}
            meta_base = {
                "source": "propagate",
                "propagation_id": op_id,
                "source_doc_id": source_doc_id,
                "seed": {"mot": mot, "marqueur": marqueur},
                "created_at": now_iso(),
            }

            for doc_id in targets:
                try:
                    self._ensure_doc(doc_id)
                    text = self._docs[doc_id].get("text", "") or ""
                    if not text:
                        results["skipped"].append({"doc_id": doc_id, "reason": "empty_text"})
                        continue

                    before = len(self._docs[doc_id]["annotations"])
                    self._annotate_similar_words(
                        doc_id=doc_id,
                        text=text,
                        mot=mot,
                        marqueur=marqueur,
                        couleur=couleur,
                        max_matches=max_matches_per_doc,
                        meta_base=meta_base,
                    )
                    after = len(self._docs[doc_id]["annotations"])
                    created = after - before
                    self._snapshot(doc_id)
                    self._docs[doc_id]["updated_at"] = now_iso()
                    results["targets"][doc_id] = {"created": created, "total": after}
                except Exception as e:
                    results["skipped"].append({"doc_id": doc_id, "reason": str(e)})

            self._persist()
            return results

    # ---- history

    def get_history(self, *, doc_id: str = _DEFAULT_DOC_ID) -> Dict[str, Any]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            doc = self._docs[doc_id]
            return {"history": list(doc["history"]), "current": int(doc["history_index"]) }

    def undo(self, *, doc_id: str = _DEFAULT_DOC_ID) -> List[Dict[str, Any]]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            doc = self._docs[doc_id]
            if doc["history_index"] <= 0:
                raise ValueError("Déjà au début de l'historique")
            doc["history_index"] -= 1
            snap = doc["history"][doc["history_index"]]
            doc["annotations"] = list(snap.get("annotations", []))
            doc["relations"] = list(snap.get("relations", []))
            self._persist()
            return list(doc["annotations"])

    def redo(self, *, doc_id: str = _DEFAULT_DOC_ID) -> List[Dict[str, Any]]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            doc = self._docs[doc_id]
            if doc["history_index"] >= len(doc["history"]) - 1:
                raise ValueError("Déjà à la fin de l'historique")
            doc["history_index"] += 1
            snap = doc["history"][doc["history_index"]]
            doc["annotations"] = list(snap.get("annotations", []))
            doc["relations"] = list(snap.get("relations", []))
            self._persist()
            return list(doc["annotations"])

    def goto(self, index: int, *, doc_id: str = _DEFAULT_DOC_ID) -> List[Dict[str, Any]]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            doc = self._docs[doc_id]
            if not (0 <= index < len(doc["history"])):
                raise ValueError("Index hors historique")
            doc["history_index"] = index
            snap = doc["history"][index]
            doc["annotations"] = list(snap.get("annotations", []))
            doc["relations"] = list(snap.get("relations", []))
            self._persist()
            return list(doc["annotations"])

    def delete_history(self, index: int, *, doc_id: str = _DEFAULT_DOC_ID) -> Dict[str, Any]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            doc = self._docs[doc_id]
            if not (0 <= index < len(doc["history"])):
                raise ValueError("Index hors historique")
            del doc["history"][index]
            if doc["history_index"] >= len(doc["history"]):
                doc["history_index"] = len(doc["history"]) - 1
            self._persist()
            return {"history": list(doc["history"]), "current": int(doc["history_index"]) }

    # ---- relations

    def list_relations(self, *, doc_id: str = _DEFAULT_DOC_ID) -> List[Dict[str, Any]]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            return list(self._docs[doc_id].get("relations", []) or [])

    def add_relation(self, rel: Dict[str, Any], *, doc_id: str = _DEFAULT_DOC_ID, dedupe: bool = True) -> Dict[str, Any]:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)

            r = dict(rel)
            r.setdefault("rel_id", new_id("rel"))
            r.setdefault("doc_id", doc_id)
            r.setdefault("created_at", now_iso())
            r.setdefault("weight", 1.0)
            r["weight"] = clamp01(float(r.get("weight", 1.0)))

            # schéma minimal
            r.setdefault("source_kind", "annotation")
            r.setdefault("target_kind", "annotation")
            r.setdefault("name", r.get("rel_type") or "relation")
            r.setdefault("rel_type", str(r.get("rel_type") or "related_to"))
            r.setdefault("meta", {})

            sk = str(r.get("source_kind") or "annotation")
            tk = str(r.get("target_kind") or "annotation")
            sid = r.get("source_id")
            tid = r.get("target_id")
            if not isinstance(sid, str) or not isinstance(tid, str) or not sid or not tid:
                raise ValueError("source_id et target_id requis")

            # validation selon le kind
            if sk == "annotation" or tk == "annotation":
                ann_by_id = self._ann_index_by_id(doc_id=doc_id)
                if sk == "annotation" and sid not in ann_by_id:
                    raise KeyError("source_id (annotation) inconnu")
                if tk == "annotation" and tid not in ann_by_id:
                    raise KeyError("target_id (annotation) inconnu")
            if sk == "marker" or tk == "marker":
                marker_names = {m.get("name") for m in (self._docs[doc_id].get("markers", []) or [])}
                if sk == "marker" and sid not in marker_names:
                    raise KeyError("source_id (marker) inconnu")
                if tk == "marker" and tid not in marker_names:
                    raise KeyError("target_id (marker) inconnu")

            # dédup
            if dedupe:
                sig = (
                    r.get("doc_id"),
                    r.get("source_kind"),
                    r.get("source_id"),
                    r.get("target_kind"),
                    r.get("target_id"),
                    r.get("rel_type"),
                    r.get("name"),
                )
                existing = {
                    (
                        x.get("doc_id"),
                        x.get("source_kind"),
                        x.get("source_id"),
                        x.get("target_kind"),
                        x.get("target_id"),
                        x.get("rel_type"),
                        x.get("name"),
                    )
                    for x in (self._docs[doc_id].get("relations", []) or [])
                }
                if sig in existing:
                    return r

            self._docs[doc_id].setdefault("relations", []).append(r)
            self._snapshot(doc_id)
            self._docs[doc_id]["updated_at"] = now_iso()
            self._persist()
            return r

    def delete_relation(self, *, doc_id: str, rel_id: str) -> None:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            rels = list(self._docs[doc_id].get("relations", []) or [])
            new_rels = [r for r in rels if r.get("rel_id") != rel_id]
            if len(new_rels) == len(rels):
                raise KeyError("Relation non trouvée")
            self._docs[doc_id]["relations"] = new_rels
            self._snapshot(doc_id)
            self._docs[doc_id]["updated_at"] = now_iso()
            self._persist()

    # ---- export / backup

    def export_csv(self, *, doc_id: str = _DEFAULT_DOC_ID) -> io.StringIO:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            anns = list(self._docs[doc_id]["annotations"])
        output = io.StringIO()
        # CSV enrichi (orienté recherche/entraînement)
        fieldnames = [
            "doc_id",
            "ann_id",
            "created_at",
            "mot",
            "marqueur",
            "couleur",
            "span_start",
            "span_end",
            "commentaire",
            "context_start",
            "context_end",
            "context",
            "phrase",
            "index_phrase",
            "meta_json",
        ]
        writer = csv.DictWriter(output, fieldnames=fieldnames)
        writer.writeheader()
        for ann in anns:
            meta = ann.get("meta", {})
            try:
                meta_json = json.dumps(meta, ensure_ascii=False)
            except Exception:
                meta_json = "{}"
            commentaire = ann.get("commentaire")
            if (not commentaire) and isinstance(meta, dict):
                commentaire = meta.get("commentaire")
            row = {
                "doc_id": ann.get("doc_id", doc_id),
                "ann_id": ann.get("ann_id", ""),
                "created_at": ann.get("created_at", ""),
                "mot": ann.get("mot", ""),
                "marqueur": ann.get("marqueur", ""),
                "couleur": ann.get("couleur", ""),
                "span_start": ann.get("span_start"),
                "span_end": ann.get("span_end"),
                "commentaire": commentaire or "",
                "context_start": ann.get("context_start"),
                "context_end": ann.get("context_end"),
                "context": ann.get("context", ""),
                "phrase": ann.get("phrase", ""),
                "index_phrase": ann.get("index_phrase", ""),
                "meta_json": meta_json,
            }
            writer.writerow(row)
        output.seek(0)
        return output

    def export_dataset_zip(self, *, doc_ids: Optional[Sequence[str]] = None) -> bytes:
        """Export "dataset" multi-tables (ZIP de CSV) adapté entraînement.

        Contient:
          - documents.csv
          - markers.csv
          - annotations.csv
          - relations.csv
        """
        with self._lock:
            if doc_ids is None:
                docs = list(self._docs.keys())
            else:
                docs = [self._validate_doc_id(d) for d in doc_ids]

            # snapshot des données nécessaires
            documents = [self.get_document_info(doc_id=d) for d in docs]
            markers_by_doc: List[Dict[str, Any]] = []
            annotations: List[Dict[str, Any]] = []
            relations: List[Dict[str, Any]] = []
            for d in docs:
                self._ensure_doc(d)
                for m in (self._docs[d].get("markers", []) or []):
                    markers_by_doc.append({"doc_id": d, "name": m.get("name"), "color": m.get("color")})
                annotations.extend(list(self._docs[d].get("annotations", []) or []))
                relations.extend(list(self._docs[d].get("relations", []) or []))

            ann_by_id: Dict[str, Dict[str, Any]] = {}
            for a in annotations:
                if isinstance(a.get("ann_id"), str):
                    ann_by_id[a["ann_id"]] = a

        # construction zip en mémoire
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            # documents.csv
            doc_csv = io.StringIO()
            doc_fields = ["doc_id", "name", "text_length", "annotation_count", "updated_at", "meta_json"]
            w = csv.DictWriter(doc_csv, fieldnames=doc_fields)
            w.writeheader()
            for d in documents:
                w.writerow(
                    {
                        "doc_id": d.get("doc_id"),
                        "name": d.get("name"),
                        "text_length": d.get("text_length"),
                        "annotation_count": d.get("annotation_count"),
                        "updated_at": d.get("updated_at"),
                        "meta_json": json.dumps(d.get("meta", {}), ensure_ascii=False),
                    }
                )
            zf.writestr("documents.csv", doc_csv.getvalue())

            # markers.csv
            m_csv = io.StringIO()
            m_fields = ["doc_id", "marker_id", "name", "color"]
            w = csv.DictWriter(m_csv, fieldnames=m_fields)
            w.writeheader()
            for m in markers_by_doc:
                doc_id = m.get("doc_id")
                name = m.get("name")
                w.writerow({"doc_id": doc_id, "marker_id": f"{doc_id}:{name}", "name": name, "color": m.get("color")})
            zf.writestr("markers.csv", m_csv.getvalue())

            # annotations.csv (réutilise le format enrichi)
            ann_csv = io.StringIO()
            ann_fields = [
                "doc_id",
                "ann_id",
                "created_at",
                "mot",
                "marqueur",
                "couleur",
                "span_start",
                "span_end",
                "context_start",
                "context_end",
                "context",
                "phrase",
                "index_phrase",
                "meta_json",
            ]
            w = csv.DictWriter(ann_csv, fieldnames=ann_fields)
            w.writeheader()
            for a in annotations:
                meta = a.get("meta", {})
                w.writerow(
                    {
                        "doc_id": a.get("doc_id"),
                        "ann_id": a.get("ann_id"),
                        "created_at": a.get("created_at"),
                        "mot": a.get("mot"),
                        "marqueur": a.get("marqueur"),
                        "couleur": a.get("couleur"),
                        "span_start": a.get("span_start"),
                        "span_end": a.get("span_end"),
                        "context_start": a.get("context_start"),
                        "context_end": a.get("context_end"),
                        "context": a.get("context"),
                        "phrase": a.get("phrase"),
                        "index_phrase": a.get("index_phrase"),
                        "meta_json": json.dumps(meta, ensure_ascii=False) if isinstance(meta, (dict, list)) else "{}",
                    }
                )
            zf.writestr("annotations.csv", ann_csv.getvalue())

            # relations.csv
            rel_csv = io.StringIO()
            rel_fields = [
                "doc_id",
                "rel_id",
                "created_at",
                "name",
                "rel_type",
                "weight",
                "source_kind",
                "source_id",
                "target_kind",
                "target_id",
                "source_mot",
                "target_mot",
                "source_marqueur",
                "target_marqueur",
                "source_span_start",
                "source_span_end",
                "target_span_start",
                "target_span_end",
                "meta_json",
            ]
            w = csv.DictWriter(rel_csv, fieldnames=rel_fields)
            w.writeheader()
            for r in relations:
                meta = r.get("meta", {})
                source = ann_by_id.get(r.get("source_id"), {}) if r.get("source_kind") == "annotation" else {}
                target = ann_by_id.get(r.get("target_id"), {}) if r.get("target_kind") == "annotation" else {}
                w.writerow(
                    {
                        "doc_id": r.get("doc_id"),
                        "rel_id": r.get("rel_id"),
                        "created_at": r.get("created_at"),
                        "name": r.get("name"),
                        "rel_type": r.get("rel_type"),
                        "weight": r.get("weight"),
                        "source_kind": r.get("source_kind"),
                        "source_id": r.get("source_id"),
                        "target_kind": r.get("target_kind"),
                        "target_id": r.get("target_id"),
                        "source_mot": source.get("mot", ""),
                        "target_mot": target.get("mot", ""),
                        "source_marqueur": source.get("marqueur", ""),
                        "target_marqueur": target.get("marqueur", ""),
                        "source_span_start": source.get("span_start"),
                        "source_span_end": source.get("span_end"),
                        "target_span_start": target.get("span_start"),
                        "target_span_end": target.get("span_end"),
                        "meta_json": json.dumps(meta, ensure_ascii=False) if isinstance(meta, (dict, list)) else "{}",
                    }
                )
            zf.writestr("relations.csv", rel_csv.getvalue())

        return buf.getvalue()

    def backup_json(self, *, doc_id: str = _DEFAULT_DOC_ID) -> io.StringIO:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            data = {
                "text": self._docs[doc_id]["text"],
                "annotations": list(self._docs[doc_id]["annotations"]),
                "relations": list(self._docs[doc_id].get("relations", []) or []),
                "markers": list(self._docs[doc_id].get("markers", []) or []),
                "version": 2,
                "doc_id": doc_id,
                "name": self._docs[doc_id].get("name", doc_id),
                "meta": self._docs[doc_id].get("meta", {}),
            }
        output = io.StringIO()
        json.dump(data, output, ensure_ascii=False)
        output.seek(0)
        return output

    def restore_json(self, data: Dict[str, Any], *, doc_id: str = _DEFAULT_DOC_ID) -> int:
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id, name=str(data.get("name") or doc_id))
            self._docs[doc_id]["text"] = str(data.get("text", "") or "")
            self._docs[doc_id]["annotations"] = list(data.get("annotations", []) or [])
            self._docs[doc_id]["relations"] = list(data.get("relations", []) or [])
            if isinstance(data.get("meta"), dict):
                self._docs[doc_id]["meta"] = dict(data.get("meta"))
            if isinstance(data.get("markers"), list):
                self._docs[doc_id]["markers"] = list(data.get("markers"))
            self._docs[doc_id]["history"] = []
            self._docs[doc_id]["history_index"] = -1
            self._snapshot(doc_id)
            self._docs[doc_id]["updated_at"] = now_iso()
            self._persist()
            return len(self._docs[doc_id]["annotations"])

    def export_xlsx_bytes(self, *, doc_id: str = _DEFAULT_DOC_ID) -> bytes:
        if Workbook is None:
            raise RuntimeError("Export XLSX indisponible: installez openpyxl")
        # Export XLSX multi-feuilles: Documents/Markers/Annotations/Relations
        with self._lock:
            doc_id = self._validate_doc_id(doc_id)
            self._ensure_doc(doc_id)
            text = self._docs[doc_id]["text"]
            anns = list(self._docs[doc_id]["annotations"])
            rels = list(self._docs[doc_id].get("relations", []) or [])
            markers = list(self._docs[doc_id].get("markers", []) or [])
            doc_info = self.get_document_info(doc_id=doc_id)
        wb = Workbook()
        ws_doc = wb.active
        ws_doc.title = "Document"
        ws_doc.append(["doc_id", "name", "text_length", "annotation_count", "updated_at", "meta_json"])
        ws_doc.append(
            [
                doc_info.get("doc_id"),
                doc_info.get("name"),
                doc_info.get("text_length"),
                doc_info.get("annotation_count"),
                doc_info.get("updated_at"),
                json.dumps(doc_info.get("meta", {}), ensure_ascii=False),
            ]
        )

        ws_m = wb.create_sheet("Markers")
        ws_m.append(["marker_id", "name", "color"])
        for m in markers:
            name = m.get("name")
            ws_m.append([name, name, m.get("color")])

        ws_a = wb.create_sheet("Annotations")
        ann_headers = [
            "doc_id",
            "ann_id",
            "created_at",
            "mot",
            "marqueur",
            "couleur",
            "span_start",
            "span_end",
            "context_start",
            "context_end",
            "context",
            "phrase",
            "index_phrase",
            "meta_json",
        ]
        ws_a.append(ann_headers)
        for a in anns:
            meta = a.get("meta", {})
            ws_a.append(
                [
                    a.get("doc_id", doc_id),
                    a.get("ann_id"),
                    a.get("created_at"),
                    a.get("mot"),
                    a.get("marqueur"),
                    a.get("couleur"),
                    a.get("span_start"),
                    a.get("span_end"),
                    a.get("context_start"),
                    a.get("context_end"),
                    a.get("context"),
                    a.get("phrase"),
                    a.get("index_phrase"),
                    json.dumps(meta, ensure_ascii=False) if isinstance(meta, (dict, list)) else "{}",
                ]
            )

        ws_r = wb.create_sheet("Relations")
        rel_headers = [
            "doc_id",
            "rel_id",
            "created_at",
            "name",
            "rel_type",
            "weight",
            "source_kind",
            "source_id",
            "target_kind",
            "target_id",
            "meta_json",
        ]
        ws_r.append(rel_headers)
        for r in rels:
            meta = r.get("meta", {})
            ws_r.append(
                [
                    r.get("doc_id", doc_id),
                    r.get("rel_id"),
                    r.get("created_at"),
                    r.get("name"),
                    r.get("rel_type"),
                    r.get("weight"),
                    r.get("source_kind"),
                    r.get("source_id"),
                    r.get("target_kind"),
                    r.get("target_id"),
                    json.dumps(meta, ensure_ascii=False) if isinstance(meta, (dict, list)) else "{}",
                ]
            )

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            tmp_path = tmp.name
        wb.save(tmp_path)
        with open(tmp_path, "rb") as f:
            data = f.read()
        try:
            os.remove(tmp_path)
        except Exception:
            pass
        return data


@lru_cache(maxsize=1)
def _get_service() -> LabellingService:
    # Persistance optionnelle (mono-utilisateur => vous pouvez la désactiver)
    # - `LABELLING_PERSIST=0|false` => aucune persistance disque
    # - sinon, `LABELLING_DATA_PATH` ou défaut local
    persist_flag = (os.environ.get("LABELLING_PERSIST", "1") or "1").strip().lower()
    persist = persist_flag not in {"0", "false", "no", "off"}
    if not persist:
        return LabellingService(data_path=None)

    data_path = os.environ.get("LABELLING_DATA_PATH")
    if data_path is None:
        data_path = os.path.join(os.path.dirname(__file__), "labelling_state.json")
    return LabellingService(data_path=data_path)


@lru_cache(maxsize=1)
def _get_ner_model():
    if pipeline is None:
        return None
    model_name = os.environ.get("LABELLING_NER_MODEL", "Gliner/gliner-ner-2")
    device = -1
    if torch is not None:
        try:
            device = 0 if torch.cuda.is_available() else -1
        except Exception:
            device = -1
    try:
        return pipeline("ner", model=model_name, device=device)
    except Exception:
        return None


# ---------------------------
# FastAPI (optionnel)
# ---------------------------

if FastAPI is not None and BaseModel is not None:
    app = FastAPI()

    # CORS optionnel (utile si le front HTML est servi depuis une autre origine)
    # Exemple:
    #   export LABELLING_CORS_ORIGINS="http://localhost:5173,http://127.0.0.1:5500"
    _cors_origins = [
        o.strip()
        for o in (os.environ.get("LABELLING_CORS_ORIGINS") or "").split(",")
        if o.strip()
    ]
    if _cors_origins and CORSMiddleware is not None:
        app.add_middleware(
            CORSMiddleware,
            allow_origins=_cors_origins,
            allow_credentials=True,
            allow_methods=["*"],
            allow_headers=["*"],
        )
else:  # pragma: no cover
    # Mode import-safe (Flask): on neutralise la partie FastAPI.
    # Les routes FastAPI ci-dessous seront décorées en no-op et ne seront pas servies.
    class _NoopFastAPI:
        def get(self, *args, **kwargs):
            return lambda fn: fn

        def post(self, *args, **kwargs):
            return lambda fn: fn

        def put(self, *args, **kwargs):
            return lambda fn: fn

        def delete(self, *args, **kwargs):
            return lambda fn: fn

    app = _NoopFastAPI()

    if BaseModel is None:
        BaseModel = object  # type: ignore[assignment]

    def _noop_param(*args, **kwargs):
        return None

    if Query is None:
        Query = _noop_param  # type: ignore[assignment]
    if Form is None:
        Form = _noop_param  # type: ignore[assignment]
    if File is None:
        File = _noop_param  # type: ignore[assignment]


# Backward-compat: MARKERS reste disponible pour du code existant.
def _markers_view() -> List[Dict[str, str]]:
    return _get_service().list_markers()


# --- Gestion dynamique des marqueurs ---
@app.post("/marker")
def add_marker(name: str = Form(...), color: str = Form(...)):
    try:
        markers = _get_service().add_marker(name=name, color=color)
        return {"status": "ok", "markers": markers}
    except ValueError as e:
        return JSONResponse(status_code=400, content={"error": str(e)})

@app.delete("/marker")
def delete_marker(name: str = Query(...)):
    markers = _get_service().delete_marker(name=name)
    return {"status": "ok", "markers": markers}

@app.put("/marker")
def update_marker(name: str = Form(...), color: str = Form(...)):
    try:
        marker = _get_service().update_marker(name=name, color=color)
        return {"status": "ok", "marker": marker}
    except KeyError as e:
        return JSONResponse(status_code=404, content={"error": str(e)})


@app.get("/markers")
def get_markers():
    return _markers_view()
# --- Recherche et filtrage d’annotations ---
@app.get("/annotations/search")
def search_annotations(
    mot: str = Query(None),
    marqueur: str = Query(None),
    phrase: str = Query(None),
    doc_id: str = Query(_DEFAULT_DOC_ID),
):
    results = _get_service().search_annotations(doc_id=doc_id, mot=mot, marqueur=marqueur, phrase=phrase)
    return {"results": results, "count": len(results)}


class DocumentCreateModel(BaseModel):
    name: str
    text: str = ""
    doc_id: Optional[str] = None
    meta: Dict[str, Any] = {}


@app.get("/documents")
def list_documents():
    return {"documents": _get_service().list_documents()}


@app.post("/documents")
def create_document(payload: DocumentCreateModel):
    try:
        doc_id = _get_service().create_document(
            name=payload.name,
            text=payload.text,
            doc_id=payload.doc_id,
            meta=payload.meta,
        )
        return {"status": "ok", "doc_id": doc_id, "document": _get_service().get_document_info(doc_id=doc_id)}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})


@app.get("/documents/{doc_id}")
def get_document(doc_id: str):
    try:
        return _get_service().get_document_info(doc_id=doc_id)
    except KeyError as e:
        return JSONResponse(status_code=404, content={"error": str(e)})
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})


@app.delete("/documents/{doc_id}")
def delete_document(doc_id: str):
    try:
        _get_service().delete_document(doc_id=doc_id)
        return {"status": "ok"}
    except KeyError as e:
        return JSONResponse(status_code=404, content={"error": str(e)})
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})


# Backward-compat (certaines fonctions historiques utilisent ces noms)
TEXT_CONTENT = ""  # maintenu, mais non-source de vérité
ANNOTATIONS: List[Dict] = []

class AnnotationModel(BaseModel):
    mot: str
    marqueur: str
    couleur: str
    phrase: str
    index_phrase: int


class PropagateRequestModel(BaseModel):
    source_doc_id: str
    target_doc_ids: List[str]
    annotation: Dict[str, Any]
    include_source_doc: bool = False
    max_matches_per_doc: int = 2000


class RelationCreateModel(BaseModel):
    doc_id: str = _DEFAULT_DOC_ID
    source_kind: str = "annotation"  # "annotation" | "marker"
    source_id: str
    target_kind: str = "annotation"  # "annotation" | "marker"
    target_id: str
    rel_type: str = "related_to"
    name: str = "relation"
    weight: float = 1.0
    meta: Dict[str, Any] = {}

def extract_text(file: Any) -> str:
    ext = file.filename.split('.')[-1].lower()
    if ext == 'txt':
        raw = file.file.read()
        try:
            return raw.decode('utf-8')
        except UnicodeDecodeError:
            return raw.decode('latin-1', errors='replace')
    elif ext == 'docx':
        if docx is None:
            raise ValueError("Support DOCX indisponible: installez python-docx")
        doc = docx.Document(file.file)
        return '\n'.join([p.text for p in doc.paragraphs])
    elif ext == 'pdf':
        if pdfplumber is None:
            raise ValueError("Support PDF indisponible: installez pdfplumber")
        with pdfplumber.open(file.file) as pdf:
            return '\n'.join([page.extract_text() or '' for page in pdf.pages])
    else:
        raise ValueError('Format de fichier non supporté')

@app.post("/upload")
async def upload_file(
    file: Any = File(...),
    doc_id: Optional[str] = Query(None),
    name: Optional[str] = Query(None),
):
    try:
        text = extract_text(file)
        svc = _get_service()
        if doc_id is None:
            # Nouveau document (multi-user / multi-session)
            doc_id = svc.create_document(
                name=name or (file.filename or "document"),
                text=text,
                meta={"filename": file.filename, "uploaded_at": now_iso()},
            )
        else:
            # Met à jour un doc existant (ou le crée implicitement)
            svc.set_text(text, doc_id=doc_id)
        if name is not None:
            svc._ensure_doc(doc_id, name=name)  # naming only
        # Backward-compat
        global TEXT_CONTENT, ANNOTATIONS
        if doc_id == _DEFAULT_DOC_ID:
            TEXT_CONTENT = text
            ANNOTATIONS = []
        return {"doc_id": doc_id, "text": text}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})


@app.post("/annotate")
def annotate(
    annotation: AnnotationModel,
    doc_id: str = Query(_DEFAULT_DOC_ID),
    auto_expand: bool = Query(True),
    dedupe: bool = Query(True),
):
    try:
        count = _get_service().add_annotation(annotation.model_dump(), doc_id=doc_id, auto_expand=auto_expand, dedupe=dedupe)
        # Backward-compat
        global ANNOTATIONS
        ANNOTATIONS = _get_service().list_annotations(doc_id=doc_id)
        return {"status": "ok", "count": count}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})


@app.post("/annotate/propagate")
def propagate_annotation(payload: PropagateRequestModel):
    try:
        res = _get_service().propagate_annotation_to_docs(
            source_doc_id=payload.source_doc_id,
            target_doc_ids=payload.target_doc_ids,
            annotation=payload.annotation,
            include_source_doc=payload.include_source_doc,
            max_matches_per_doc=payload.max_matches_per_doc,
        )
        return {"status": "ok", **res}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})


@app.get("/relations")
def list_relations(doc_id: str = Query(_DEFAULT_DOC_ID)):
    try:
        return {"relations": _get_service().list_relations(doc_id=doc_id)}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})


@app.post("/relations")
def create_relation(payload: RelationCreateModel, dedupe: bool = Query(True)):
    try:
        rel = _get_service().add_relation(payload.model_dump(), doc_id=payload.doc_id, dedupe=dedupe)
        return {"status": "ok", "relation": rel}
    except KeyError as e:
        return JSONResponse(status_code=404, content={"error": str(e)})
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})


@app.delete("/relations/{rel_id}")
def delete_relation(rel_id: str, doc_id: str = Query(_DEFAULT_DOC_ID)):
    try:
        _get_service().delete_relation(doc_id=doc_id, rel_id=rel_id)
        return {"status": "ok"}
    except KeyError as e:
        return JSONResponse(status_code=404, content={"error": str(e)})
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})
# --- Endpoints pour l'historique ---

@app.get("/history")
def get_history(doc_id: str = Query(_DEFAULT_DOC_ID)):
    return _get_service().get_history(doc_id=doc_id)

@app.post("/history/undo")
def undo(doc_id: str = Query(_DEFAULT_DOC_ID)):
    try:
        anns = _get_service().undo(doc_id=doc_id)
        global ANNOTATIONS
        if doc_id == _DEFAULT_DOC_ID:
            ANNOTATIONS = anns
        h = _get_service().get_history(doc_id=doc_id)
        return {"status": "ok", "current": h["current"], "annotations": anns}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})

@app.post("/history/redo")
def redo(doc_id: str = Query(_DEFAULT_DOC_ID)):
    try:
        anns = _get_service().redo(doc_id=doc_id)
        global ANNOTATIONS
        if doc_id == _DEFAULT_DOC_ID:
            ANNOTATIONS = anns
        h = _get_service().get_history(doc_id=doc_id)
        return {"status": "ok", "current": h["current"], "annotations": anns}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})

@app.post("/history/goto")
def goto(index: int = Form(...), doc_id: str = Form(_DEFAULT_DOC_ID)):
    try:
        anns = _get_service().goto(index, doc_id=doc_id)
        global ANNOTATIONS
        if doc_id == _DEFAULT_DOC_ID:
            ANNOTATIONS = anns
        h = _get_service().get_history(doc_id=doc_id)
        return {"status": "ok", "current": h["current"], "annotations": anns}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})

@app.delete("/history/delete")
def delete_history(index: int = Form(...), doc_id: str = Form(_DEFAULT_DOC_ID)):
    try:
        return {"status": "ok", **_get_service().delete_history(index, doc_id=doc_id)}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})

@app.get("/csv")
def get_csv(doc_id: str = Query(_DEFAULT_DOC_ID)):
    output = _get_service().export_csv(doc_id=doc_id)
    return StreamingResponse(output, media_type="text/csv", headers={"Content-Disposition": "attachment; filename=annotations.csv"})


@app.get("/dataset.zip")
def download_dataset_zip(doc_id: Optional[str] = Query(None)):
    try:
        if doc_id is None:
            data = _get_service().export_dataset_zip(doc_ids=None)
        else:
            data = _get_service().export_dataset_zip(doc_ids=[doc_id])
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=dataset.zip"},
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})

@app.get("/backup")
def get_backup(doc_id: str = Query(_DEFAULT_DOC_ID)):
    output = _get_service().backup_json(doc_id=doc_id)
    return StreamingResponse(output, media_type="application/json", headers={"Content-Disposition": "attachment; filename=backup.json"})

@app.post("/restore")
async def restore(file: Any = File(...), doc_id: str = Query(_DEFAULT_DOC_ID)):
    try:
        data = json.load(file.file)
        count = _get_service().restore_json(data, doc_id=doc_id)
        # Backward-compat
        global TEXT_CONTENT, ANNOTATIONS
        if doc_id == _DEFAULT_DOC_ID:
            TEXT_CONTENT = _get_service().get_text(doc_id=_DEFAULT_DOC_ID)
            ANNOTATIONS = _get_service().list_annotations(doc_id=_DEFAULT_DOC_ID)
            return {"status": "ok", "doc_id": doc_id, "text": TEXT_CONTENT, "count": count}
        return {"status": "ok", "doc_id": doc_id, "count": count}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})


@app.get("/xlsx")
def download_xlsx(doc_id: str = Query(_DEFAULT_DOC_ID)):
    try:
        data = _get_service().export_xlsx_bytes(doc_id=doc_id)
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=annotations.xlsx"},
        )
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})


@app.get("/suggestions")
def get_suggestions(doc_id: str = Query(_DEFAULT_DOC_ID)):
    model = _get_ner_model()
    if model is None:
        return JSONResponse(status_code=503, content={"error": "Modèle NER non disponible (installez transformers/torch)"})

    text = _get_service().get_text(doc_id=doc_id)
    if not text:
        return {"suggestions": []}

    results = model(text)
    suggestions: List[Dict[str, Any]] = []
    for ent in results:
        score = float(ent.get("score", 1.0))
        highlight = "yellow" if score < 0.7 else "none"
        suggestions.append(
            {
                "mot": ent.get("word", ""),
                "marqueur": ent.get("entity_group", ""),
                "couleur": "#FFFF00" if highlight == "yellow" else "#00FF00",
                "phrase": ent.get("sentence", ""),
                "index_phrase": 0,
                "score": score,
                "highlight": highlight,
            }
        )
    return {"suggestions": suggestions}


def create_flask_blueprint(*, url_prefix: str = "/labelling"):
    """Point d’extension Flask.

    Ne requiert pas Flask au runtime tant que la fonction n’est pas appelée.
    Usage:
        from annotation import create_flask_blueprint
        app.register_blueprint(create_flask_blueprint(), url_prefix="/api")
    """

    try:
        from flask import Blueprint, jsonify, make_response, request  # type: ignore
    except Exception as e:  # pragma: no cover
        raise RuntimeError("Flask n'est pas installé") from e

    bp = Blueprint("labelling", __name__)

    @bp.get(f"{url_prefix}/markers")
    def _flask_markers():
        doc_id = (request.args.get("doc_id") or _DEFAULT_DOC_ID).strip()
        return jsonify(_get_service().list_markers(doc_id=doc_id))

    @bp.post(f"{url_prefix}/markers")
    def _flask_add_marker():
        payload = request.get_json(silent=True) or {}
        name = str(payload.get("name") or "").strip()
        color = str(payload.get("color") or "").strip()
        doc_id = str(payload.get("doc_id") or request.args.get("doc_id") or _DEFAULT_DOC_ID).strip()
        if not name or not color:
            return jsonify({"error": "name et color requis"}), 400
        try:
            markers = _get_service().add_marker(name=name, color=color, doc_id=doc_id)
            return jsonify({"status": "ok", "markers": markers})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.put(f"{url_prefix}/markers/<name>")
    def _flask_update_marker(name: str):
        payload = request.get_json(silent=True) or {}
        color = str(payload.get("color") or "").strip()
        doc_id = (request.args.get("doc_id") or _DEFAULT_DOC_ID).strip()
        if not color:
            return jsonify({"error": "color requis"}), 400
        try:
            marker = _get_service().update_marker(name=name, color=color, doc_id=doc_id)
            return jsonify({"status": "ok", "marker": marker})
        except KeyError as ex:
            return jsonify({"error": str(ex)}), 404
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.delete(f"{url_prefix}/markers/<name>")
    def _flask_delete_marker(name: str):
        try:
            doc_id = (request.args.get("doc_id") or _DEFAULT_DOC_ID).strip()
            markers = _get_service().delete_marker(name=name, doc_id=doc_id)
            return jsonify({"status": "ok", "markers": markers})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.get(f"{url_prefix}/documents")
    def _flask_list_documents():
        return jsonify({"documents": _get_service().list_documents()})

    @bp.get(f"{url_prefix}/documents/<doc_id>")
    def _flask_get_document(doc_id: str):
        try:
            svc = _get_service()
            return jsonify(
                {
                    "info": svc.get_document_info(doc_id=doc_id),
                    "text": svc.get_text(doc_id=doc_id),
                    "annotations": svc.list_annotations(doc_id=doc_id),
                    "relations": svc.list_relations(doc_id=doc_id),
                    "history": svc.get_history(doc_id=doc_id),
                }
            )
        except KeyError as ex:
            return jsonify({"error": str(ex)}), 404
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.delete(f"{url_prefix}/documents/<doc_id>")
    def _flask_delete_document(doc_id: str):
        try:
            _get_service().delete_document(doc_id=doc_id)
            return jsonify({"status": "ok"})
        except KeyError as ex:
            return jsonify({"error": str(ex)}), 404
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.post(f"{url_prefix}/documents/<doc_id>/text")
    def _flask_set_document_text(doc_id: str):
        payload = request.get_json(silent=True) or {}
        text = str(payload.get("text") or "")
        try:
            _get_service().set_text(text, doc_id=doc_id)
            return jsonify({"status": "ok", "doc_id": doc_id})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.post(f"{url_prefix}/documents")
    def _flask_create_document():
        payload = request.get_json(silent=True) or {}
        try:
            doc_id = _get_service().create_document(
                name=str(payload.get("name") or "document"),
                text=str(payload.get("text") or ""),
                doc_id=payload.get("doc_id"),
                meta=payload.get("meta") if isinstance(payload.get("meta"), dict) else {},
            )
            return jsonify({"status": "ok", "doc_id": doc_id})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.post(f"{url_prefix}/upload")
    def _flask_upload():
        if "file" not in request.files:
            return jsonify({"error": "file manquant"}), 400
        f = request.files["file"]
        doc_id = request.args.get("doc_id")
        name = request.args.get("name")
        # On réutilise la logique d'extraction via un petit adaptateur minimal.
        class _UF:
            filename = f.filename
            file = f.stream

        try:
            text = extract_text(_UF)  # type: ignore[arg-type]
            svc = _get_service()
            if not doc_id:
                doc_id = svc.create_document(
                    name=name or (f.filename or "document"),
                    text=text,
                    meta={"filename": f.filename, "uploaded_at": now_iso()},
                )
            else:
                svc.set_text(text, doc_id=doc_id)
                if name:
                    svc._ensure_doc(doc_id, name=name)  # naming only
            return jsonify({"doc_id": doc_id, "text": text})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.post(f"{url_prefix}/annotate")
    def _flask_annotate():
        payload = request.get_json(silent=True) or {}
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        try:
            count = _get_service().add_annotation(payload, doc_id=doc_id)
            return jsonify({"status": "ok", "count": count})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.get(f"{url_prefix}/annotations")
    def _flask_list_annotations():
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        try:
            return jsonify({"annotations": _get_service().list_annotations(doc_id=doc_id)})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.post(f"{url_prefix}/annotations/delete_by_span")
    def _flask_delete_annotations_by_span():
        payload = request.get_json(silent=True) or {}
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        try:
            s0 = int(payload.get("span_start"))
            s1 = int(payload.get("span_end"))
        except Exception:
            return jsonify({"error": "span_start/span_end requis"}), 400
        try:
            res = _get_service().delete_annotations_by_span(doc_id=doc_id, span_start=s0, span_end=s1)
            return jsonify({"status": "ok", **res})
        except KeyError as ex:
            return jsonify({"error": str(ex)}), 404
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.get(f"{url_prefix}/annotations/search")
    def _flask_search_annotations():
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        mot = request.args.get("mot")
        marqueur = request.args.get("marqueur")
        phrase = request.args.get("phrase")
        try:
            results = _get_service().search_annotations(doc_id=doc_id, mot=mot, marqueur=marqueur, phrase=phrase)
            return jsonify({"results": results, "count": len(results)})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.post(f"{url_prefix}/annotate/propagate")
    def _flask_propagate():
        payload = request.get_json(silent=True) or {}
        try:
            res = _get_service().propagate_annotation_to_docs(
                source_doc_id=str(payload.get("source_doc_id") or _DEFAULT_DOC_ID),
                target_doc_ids=payload.get("target_doc_ids") or [],
                annotation=payload.get("annotation") or {},
                include_source_doc=bool(payload.get("include_source_doc") or False),
                max_matches_per_doc=int(payload.get("max_matches_per_doc") or 2000),
            )
            return jsonify({"status": "ok", **res})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.get(f"{url_prefix}/relations")
    def _flask_list_relations():
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        try:
            return jsonify({"relations": _get_service().list_relations(doc_id=doc_id)})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.post(f"{url_prefix}/relations")
    def _flask_create_relation():
        payload = request.get_json(silent=True) or {}
        doc_id = payload.get("doc_id") or (request.args.get("doc_id") or _DEFAULT_DOC_ID)
        try:
            rel = _get_service().add_relation(payload, doc_id=doc_id, dedupe=True)
            return jsonify({"status": "ok", "relation": rel})
        except KeyError as ex:
            return jsonify({"error": str(ex)}), 404
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.delete(f"{url_prefix}/relations/<rel_id>")
    def _flask_delete_relation(rel_id: str):
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        try:
            _get_service().delete_relation(doc_id=doc_id, rel_id=rel_id)
            return jsonify({"status": "ok"})
        except KeyError as ex:
            return jsonify({"error": str(ex)}), 404
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    # ---- Historique

    @bp.get(f"{url_prefix}/history")
    def _flask_get_history():
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        try:
            return jsonify(_get_service().get_history(doc_id=doc_id))
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.post(f"{url_prefix}/history/undo")
    def _flask_undo():
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        try:
            anns = _get_service().undo(doc_id=doc_id)
            h = _get_service().get_history(doc_id=doc_id)
            return jsonify({"status": "ok", "current": h.get("current"), "annotations": anns})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.post(f"{url_prefix}/history/redo")
    def _flask_redo():
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        try:
            anns = _get_service().redo(doc_id=doc_id)
            h = _get_service().get_history(doc_id=doc_id)
            return jsonify({"status": "ok", "current": h.get("current"), "annotations": anns})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.post(f"{url_prefix}/history/goto")
    def _flask_goto():
        payload = request.get_json(silent=True) or {}
        doc_id = str(payload.get("doc_id") or _DEFAULT_DOC_ID)
        try:
            index = int(payload.get("index"))
        except Exception:
            return jsonify({"error": "index invalide"}), 400
        try:
            anns = _get_service().goto(index, doc_id=doc_id)
            h = _get_service().get_history(doc_id=doc_id)
            return jsonify({"status": "ok", "current": h.get("current"), "annotations": anns})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.delete(f"{url_prefix}/history/<int:index>")
    def _flask_delete_history(index: int):
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        try:
            return jsonify({"status": "ok", **_get_service().delete_history(index, doc_id=doc_id)})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    # ---- Exports / backup

    @bp.get(f"{url_prefix}/export/csv")
    def _flask_export_csv():
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        try:
            output = _get_service().export_csv(doc_id=doc_id)
            resp = make_response(output.getvalue())
            resp.headers["Content-Type"] = "text/csv; charset=utf-8"
            safe = re.sub(r"[^a-zA-Z0-9._-]+", "_", str(doc_id)) or "document"
            resp.headers["Content-Disposition"] = f"attachment; filename={safe}_annotations.csv"
            return resp
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.get(f"{url_prefix}/export/csv_view")
    def _flask_export_csv_view():
        """Export CSV orienté UI (mêmes colonnes que le tableau de la page annotate).

                Colonnes:
                    - element_surligne
                    - phrase_contexte
                    - marqueur1..3 (max 3 marqueurs par span)
                    - commentaire (agrégé par span)
        """

        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        try:
            svc = _get_service()
            text = svc.get_text(doc_id=doc_id) or ""
            anns = svc.list_annotations(doc_id=doc_id) or []

            by_span: Dict[tuple, Dict[str, Any]] = {}
            for a in anns:
                try:
                    start = int(a.get("span_start"))
                    end = int(a.get("span_end"))
                except Exception:
                    continue
                if not (0 <= start < end <= len(text)):
                    continue

                key = (start, end)
                if key not in by_span:
                    by_span[key] = {
                        "start": start,
                        "end": end,
                        "element": (a.get("mot") or text[start:end] or ""),
                        "phrase": (a.get("phrase") or a.get("context") or ""),
                        "markers": [],
                        "comments": [],
                    }

                m = str(a.get("marqueur") or "").strip()
                if m and m not in by_span[key]["markers"]:
                    by_span[key]["markers"].append(m)

                meta = a.get("meta") if isinstance(a.get("meta"), dict) else {}
                c = str((a.get("commentaire") or meta.get("commentaire") or "")).strip()
                if c and c not in by_span[key]["comments"]:
                    by_span[key]["comments"].append(c)

            rows = sorted(by_span.values(), key=lambda r: (r.get("start", 0), r.get("end", 0)))

            output = io.StringIO()
            fieldnames = ["element_surligne", "phrase_contexte", "marqueur1", "marqueur2", "marqueur3", "commentaire"]
            w = csv.DictWriter(output, fieldnames=fieldnames)
            w.writeheader()
            for r in rows:
                ms = (r.get("markers") or [])[:3]
                comment = " | ".join((r.get("comments") or [])[:50])
                w.writerow(
                    {
                        "element_surligne": r.get("element", "") or "",
                        "phrase_contexte": r.get("phrase", "") or "",
                        "marqueur1": ms[0] if len(ms) > 0 else "",
                        "marqueur2": ms[1] if len(ms) > 1 else "",
                        "marqueur3": ms[2] if len(ms) > 2 else "",
                        "commentaire": comment,
                    }
                )

            output.seek(0)
            resp = make_response(output.getvalue())
            resp.headers["Content-Type"] = "text/csv; charset=utf-8"
            safe = re.sub(r"[^a-zA-Z0-9._-]+", "_", str(doc_id)) or "document"
            resp.headers["Content-Disposition"] = f"attachment; filename={safe}_annotations_view.csv"
            return resp
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.get(f"{url_prefix}/export/xlsx")
    def _flask_export_xlsx():
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        try:
            data = _get_service().export_xlsx_bytes(doc_id=doc_id)
            resp = make_response(data)
            resp.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            safe = re.sub(r"[^a-zA-Z0-9._-]+", "_", str(doc_id)) or "document"
            resp.headers["Content-Disposition"] = f"attachment; filename={safe}_annotations.xlsx"
            return resp
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.get(f"{url_prefix}/backup")
    def _flask_backup():
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        try:
            output = _get_service().backup_json(doc_id=doc_id)
            resp = make_response(output.getvalue())
            resp.headers["Content-Type"] = "application/json; charset=utf-8"
            resp.headers["Content-Disposition"] = "attachment; filename=backup.json"
            return resp
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.post(f"{url_prefix}/restore")
    def _flask_restore():
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        try:
            data = None

            if request.files and "file" in request.files:
                f = request.files["file"]
                data = json.load(f.stream)
            else:
                payload = request.get_json(silent=True)
                if isinstance(payload, dict) and "data" in payload and isinstance(payload["data"], dict):
                    data = payload["data"]
                elif isinstance(payload, dict):
                    data = payload

            if not isinstance(data, dict):
                return jsonify({"error": "Backup JSON invalide"}), 400

            count = _get_service().restore_json(data, doc_id=doc_id)
            return jsonify({"status": "ok", "doc_id": doc_id, "count": count})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.get(f"{url_prefix}/suggestions")
    def _flask_suggestions():
        doc_id = request.args.get("doc_id") or _DEFAULT_DOC_ID
        model = _get_ner_model()
        if model is None:
            return jsonify({"error": "Modèle NER non disponible (installez transformers/torch)"}), 503

        try:
            text = _get_service().get_text(doc_id=doc_id)
            if not text:
                return jsonify({"suggestions": []})

            results = model(text)
            suggestions: List[Dict[str, Any]] = []
            for ent in results:
                score = float(ent.get("score", 1.0))
                highlight = "yellow" if score < 0.7 else "none"
                suggestions.append(
                    {
                        "mot": ent.get("word", ""),
                        "marqueur": ent.get("entity_group", ""),
                        "couleur": "#FFFF00" if highlight == "yellow" else "#00FF00",
                        "phrase": ent.get("sentence", ""),
                        "index_phrase": 0,
                        "score": score,
                        "highlight": highlight,
                    }
                )
            return jsonify({"suggestions": suggestions})
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    @bp.get(f"{url_prefix}/dataset.zip")
    def _flask_dataset_zip():
        doc_id = request.args.get("doc_id")
        try:
            data = _get_service().export_dataset_zip(doc_ids=[doc_id] if doc_id else None)
            resp = make_response(data)
            resp.headers["Content-Type"] = "application/zip"
            resp.headers["Content-Disposition"] = "attachment; filename=dataset.zip"
            return resp
        except Exception as ex:
            return jsonify({"error": str(ex)}), 400

    return bp


# ---------------------------
# Excel export (annotations + relations)
# ---------------------------

class ExcelExporter:
    """
    Exporte 2 feuilles:
      - Annotations: doc, ann_id, labels, weight, weight_norm, text, sentence, span, meta
      - Relations: doc, rel_id, type, weight, source_ann_id, target_ann_id, meta
    """

    def __init__(self, store: AnnotationStore) -> None:
        self.store = store

    def export(
        self,
        filepath: str,
        *,
        doc_id: str,
        normalize_weights: str = "identity",
        include_meta: bool = True,
    ) -> str:
        doc = self.store.get_document(doc_id)
        anns = self.store.list_annotations(doc_id=doc_id)
        rels = self.store.list_relations(doc_id=doc_id)

        weights = [a.weight for a in anns]
        w_norm = WeightNormalizer.normalize(weights, policy=normalize_weights)
        ann_rows = list(zip(anns, w_norm))

        wb = Workbook()
        wsA = wb.active
        wsA.title = "Annotations"
        wsR = wb.create_sheet("Relations")

        self._write_annotations_sheet(wsA, doc, ann_rows, include_meta=include_meta)
        self._write_relations_sheet(wsR, doc, rels, include_meta=include_meta)

        wb.save(filepath)
        return filepath

    def _write_annotations_sheet(self, ws, doc: Document, ann_rows, *, include_meta: bool) -> None:
        headers = [
            "doc_id", "doc_name", "ann_id",
            "labels", "labels_names",
            "weight", "weight_norm",
            "text", "sentence",
            "span_start", "span_end",
            "created_at",
        ]
        if include_meta:
            headers.append("meta_json")

        ws.append(headers)
        self._style_header(ws, len(headers))

        for ann, wn in ann_rows:
            label_names = [self.store.get_label(lid).name for lid in ann.labels]
            row = [
                doc.id, doc.name, ann.id,
                ",".join(ann.labels),
                ",".join(label_names),
                float(ann.weight),
                float(wn),
                ann.text,
                ann.sentence,
                int(ann.span.start),
                int(ann.span.end),
                ann.created_at,
            ]
            if include_meta:
                row.append(str(ann.meta))
            ws.append(row)

        self._autosize(ws)

    def _write_relations_sheet(self, ws, doc: Document, rels: List[Relation], *, include_meta: bool) -> None:
        headers = [
            "doc_id", "doc_name", "rel_id",
            "rel_type", "weight",
            "source_ann_id", "target_ann_id",
            "created_at",
        ]
        if include_meta:
            headers.append("meta_json")

        ws.append(headers)
        self._style_header(ws, len(headers))

        for r in rels:
            row = [
                doc.id, doc.name, r.id,
                r.rel_type, float(r.weight),
                r.source_ann_id, r.target_ann_id,
                r.created_at,
            ]
            if include_meta:
                row.append(str(r.meta))
            ws.append(row)

        self._autosize(ws)

    @staticmethod
    def _style_header(ws, ncols: int) -> None:
        fill = PatternFill("solid", fgColor="1F2937")  # dark header
        font = Font(color="FFFFFF", bold=True)
        align = Alignment(vertical="center", horizontal="center", wrap_text=True)
        for c in range(1, ncols + 1):
            cell = ws.cell(row=1, column=c)
            cell.fill = fill
            cell.font = font
            cell.alignment = align
        ws.freeze_panes = "A2"

    @staticmethod
    def _autosize(ws) -> None:
        for col in range(1, ws.max_column + 1):
            maxlen = 0
            for row in range(1, ws.max_row + 1):
                v = ws.cell(row=row, column=col).value
                if v is None:
                    continue
                maxlen = max(maxlen, len(str(v)))
            ws.column_dimensions[get_column_letter(col)].width = min(60, max(10, maxlen + 2))




# À ajouter dans ton backend d’annotation (ex: annotation_store.py)


def regex_annotate(
    store,
    *,
    doc_id: str,
    pattern: str,
    label_ids: Iterable[str],
    weight: float = 1.0,
    flags: int = re.IGNORECASE,
    allow_overlap: bool = False,
    group: int = 0,
    meta: Optional[Dict[str, Any]] = None,
):
    """
    Crée automatiquement des annotations à partir d’une recherche regex.

    Paramètres
    ----------
    store : AnnotationStore
        Instance de ton store Quantix.
    doc_id : str
        ID du document à annoter.
    pattern : str
        Expression régulière.
    label_ids : Iterable[str]
        Liste d'IDs de labels à appliquer (multi-label supporté).
    weight : float, default=1.0
        Poids de l’annotation (sera clampé dans [0,1]).
    flags : int, default=re.IGNORECASE
        Flags regex.
    allow_overlap : bool, default=False
        Si False, empêche les annotations qui se recouvrent.
    group : int, default=0
        Groupe regex à utiliser pour le span (0 = match complet).
    meta : dict, optional
        Métadonnées supplémentaires.

    Retour
    ------
    List[Annotation]
        Liste des annotations créées.
    """

    doc = store.get_document(doc_id)
    text = doc.raw_text

    try:
        regex: Pattern = re.compile(pattern, flags)
    except re.error as e:
        raise ValueError(f"Invalid regex: {e}")

    created = []
    occupied = []  # intervalles déjà annotés si allow_overlap=False

    for match in regex.finditer(text):
        try:
            start, end = match.span(group)
        except IndexError:
            raise ValueError(f"Invalid group={group} for pattern")

        if start == end:
            continue

        if not allow_overlap:
            overlap = any(not (end <= s or start >= e) for s, e in occupied)
            if overlap:
                continue

        ann = store.add_annotation(
            doc_id=doc_id,
            start=start,
            end=end,
            label_ids=list(label_ids),
            weight=weight,
            meta={
                "source": "regex",
                "pattern": pattern,
                "group": group,
                **(meta or {}),
            },
        )

        created.append(ann)
        occupied.append((start, end))

    return created


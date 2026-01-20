import shutil
from pathlib import Path
from copy import deepcopy
from .text_translation import translate_text
from .math_detection import is_mathematical_content


def _is_protected_run(run):
    el = run._element
    if el.xpath(".//m:oMath"):
        return True
    if el.xpath(".//w:fldChar"):
        return True
    if el.xpath(".//w:hyperlink"):
        return True
    return False


def _group_runs_by_style(paragraph):
    groups = []
    current = []

    def style_signature(r):
        f = r.font
        return (
            r.bold, r.italic, r.underline,
            f.strike, f.double_strike, f.all_caps, f.small_caps,
            f.name, f.size,
            f.color.rgb if f.color else None,
            f.highlight_color,
            f.superscript, f.subscript,
            f.emboss, f.imprint, f.outline, f.shadow
        )

    prev_sig = None
    for run in paragraph.runs:
        if _is_protected_run(run):
            if current:
                groups.append(current)
                current = []
            groups.append([run])
            prev_sig = None
            continue

        sig = style_signature(run)
        if sig != prev_sig and current:
            groups.append(current)
            current = []
        current.append(run)
        prev_sig = sig

    if current:
        groups.append(current)
    return groups


def _translate_paragraph(paragraph, translator):
    if not paragraph.text.strip():
        return
    if is_mathematical_content(paragraph.text):
        return

    groups = _group_runs_by_style(paragraph)

    for grp in groups:
        if len(grp) == 1 and _is_protected_run(grp[0]):
            continue

        text = "".join(r.text for r in grp)
        if not text.strip():
            continue

        translated = translate_text(text, translator)

        grp[0].text = translated
        for r in grp[1:]:
            r.text = ""


def translate_docx(input_path, output_path, translator):
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise ImportError(
            "Le paquet 'python-docx' est requis pour traduire les fichiers .docx. "
            "Installez-le avec `pip install python-docx`."
        ) from e

    input_path = Path(input_path)
    output_path = Path(output_path)

    shutil.copy2(input_path, output_path)
    doc = Document(output_path)

    for para in doc.paragraphs:
        _translate_paragraph(para, translator)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _translate_paragraph(para, translator)

    for section in doc.sections:
        for hf in (section.header, section.footer):
            for para in hf.paragraphs:
                _translate_paragraph(para, translator)

    doc.save(output_path)
